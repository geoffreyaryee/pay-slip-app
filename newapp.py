import pandas as pd
from docx import Document
import comtypes.client  # For Word to PDF conversion
from tkinter import Tk, Label, Button, filedialog, messagebox
import os
import threading

# Ghana PAYE tax calculation function
def calculate_paye(taxable):
    tax = 0
    bands = [
        (490, 0.0),  # First GHS 528 at 0%
        (110, 0.05),  # Next GHS 100 at 5%
        (130, 0.10),  # Next GHS 160 at 10%
        (3166.67, 0.175),  # Next GHS 3020 at 17.5%
        (16302, 0.25),  # Next GHS 16,302 at 25%
        (float('inf'), 0.30),  # Amount above GHS 20,110 at 30%
    ]

    remaining_income = taxable
    for band_limit, rate in bands:
        if remaining_income > band_limit:
            tax += band_limit * rate
            remaining_income -= band_limit
        else:
            tax += remaining_income * rate
            break
    return tax

# Function to perform gross pay, PAYE tax, and net pay calculations
def calculate_net_pay(row):
    gross_pay = row['Basic Salary'] + row['Allowances']
    ssnit = gross_pay * 0.055
    taxable_income = gross_pay - ssnit
    paye = calculate_paye(taxable_income)  # Calculate PAYE tax
    net_pay = taxable_income - paye - row['Deductions']
    return gross_pay, paye, net_pay, ssnit, taxable_income

# Function to generate payslip from a Word template
def generate_payslip_from_template(employee_data, template_path, output_dir):
    # Load the Word template
    doc = Document(template_path)

    # Replace placeholders with employee data
    for paragraph in doc.paragraphs:
        for key, value in employee_data.items():
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Save the generated payslip as a new Word document
    output_path = os.path.join(output_dir, f"payslip_{employee_data['Name']}.docx")
    doc.save(output_path)
    print(f"Payslip generated: {output_path}")

    # Convert to PDF
    convert_to_pdf(output_path)

# Function to convert Word document to PDF
def convert_to_pdf(docx_path):
    try:
        # Normalize the path to avoid issues with backslashes or spaces
        docx_path = os.path.abspath(docx_path)
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'

        # Use Word COM interface to convert
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # FileFormat 17 = PDF
        doc.Close()
        word.Quit()

        print(f"Converted to PDF: {pdf_path}")
    except Exception as e:
        print(f"Failed to convert {docx_path} to PDF: {str(e)}")

# Function to load the Excel file and generate payslips
def load_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", ".xlsx;.xls")])
    if file_path:
        try:
            # Load the Excel data
            df = pd.read_excel(file_path)

            # Validate required columns
            required_columns = ['Name', 'Basic Salary', 'Allowances', 'Deductions']
            for col in required_columns:
                if col not in df.columns:
                    raise ValueError(f"Missing required column: {col}")

            # Calculate Gross Pay, PAYE, and Net Pay
            df['Gross Pay'], df['PAYE'], df['Net Pay'], df['SSNIT'], df['Taxable'] = zip(
                *df.apply(calculate_net_pay, axis=1)
            )

            # Select the Word template for generating payslips
            template_path = filedialog.askopenfilename(filetypes=[("Word files", ".docx")])
            if not template_path:
                messagebox.showwarning("Warning", "No template selected!")
                return

            # Select output directory
            output_dir = filedialog.askdirectory()
            if not output_dir:
                messagebox.showwarning("Warning", "No output directory selected!")
                return

            # Validate placeholders in template
            doc = Document(template_path)
            template_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            required_placeholders = [
                '{Name}', '{Basic Salary}', '{Allowances}', '{Gross Pay}', '{PAYE}',
                '{Deductions}', '{Net Pay}', '{SSNIT}', '{Taxable}', '{Year}', '{Role}', '{Month}'
            ]
            missing_placeholders = [
                ph for ph in required_placeholders if ph not in template_content
            ]
            if missing_placeholders:
                messagebox.showerror(
                    "Error", f"Missing placeholders in template: {', '.join(missing_placeholders)}"
                )
                return

            # Generate payslips in a separate thread to keep UI responsive
            def generate_payslips():
                for _, row in df.iterrows():
                    employee_data = row.to_dict()
                    generate_payslip_from_template(employee_data, template_path, output_dir)
                messagebox.showinfo("Success", "Payslips generated successfully!")

            threading.Thread(target=generate_payslips).start()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to process file: {str(e)}")
    else:
        messagebox.showwarning("Warning", "No file selected!")

# GUI Setup
root = Tk()
root.title("Pay Slip Generator")

label = Label(root, text="Select Excel File to Generate Pay Slips")
label.pack(pady=20)

button = Button(root, text="Load Excel File", command=load_file)
button.pack(pady=10)

root.mainloop()
